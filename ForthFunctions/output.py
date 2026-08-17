import sys
from datetime import datetime
from pathlib import Path

from docx import Document


def output_playlists(statrequest, keywords=None):
    if keywords:
        print(f"Your request: {keywords}")

    parsed_playlists = []

    for number, item in enumerate(statrequest["items"], start=1):
        title = item["snippet"]["title"]
        playlist_id = item["id"]
        published_at = item["snippet"]["publishedAt"]
        dt = datetime.fromisoformat(published_at.replace("Z", "+00:00"))
        formatted_date = dt.strftime("%d.%m.%Y %H:%M:%S")

        status = item["status"]["privacyStatus"]
        itemCount = item["contentDetails"]["itemCount"]

        channelName = item["snippet"]["channelTitle"]
        channelId = item["snippet"]["channelId"]

        parsed_playlists.append({
            "title": title,
            "playlist_id": playlist_id,
            "formatted_date": formatted_date,
            "status": status,
            "itemCount": itemCount,
            "channelName": channelName,
            "channelId": channelId
        })

        print(
            f"\n{number}."
            f"\n{title}\n"
            f"Playlist URL: https://www.youtube.com/playlist?list={playlist_id}\n"
            f"Privacy status: {status}; Videos: {itemCount}\n"
            f"{formatted_date}\n"
            f"{channelName}\n"
            f"Channel URL: https://www.youtube.com/channel/{channelId}")

    return parsed_playlists



def save_docx(parsed_playlists, keywords=None):
    choice = input("\n\nDo you want to save videos in a DOCX file? (y/n): ").strip().lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").strip().lower()
        
    if choice == "n":
        return choice, None

    elif choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = Path(app_folder, "Playlists_in_DOCX")
        youtube_folder.mkdir(parents=True, exist_ok=True)

        full_path = Path(youtube_folder, f"{keywords}.docx")

        counter = 0
        while full_path.exists():
            counter += 1
            full_path = Path(youtube_folder, f"{keywords} ({counter}).docx")

        doc = Document()
        doc.add_heading("Playlist Report for YouTube", 0)
        
        if keywords:
            doc.add_heading(f"Request: {keywords}", level=1)
            
        doc.add_heading("The list of playlists:", level=1)
        
        for number, playlist in enumerate(parsed_playlists, start=1):
            doc.add_heading(f"{number}. {playlist['title']}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Playlist Link: ").bold = True
            p.add_run(f"https://www.youtube.com/playlist?list={playlist['playlist_id']}\n")
            
            p.add_run("Privacy status: ").bold = True
            p.add_run(f"{playlist['status']}; ")
            
            p.add_run("Videos: ").bold = True
            p.add_run(f"{playlist['itemCount']}\n")
            
            p.add_run("Date: ").bold = True
            p.add_run(f"{playlist['formatted_date']}\n")
            
            p.add_run("Channel: ").bold = True
            p.add_run(f"{playlist['channelName']}\n")
            
            p.add_run("Channel URL: ").bold = True
            p.add_run(f"https://www.youtube.com/channel/{playlist['channelId']}")
            
        doc.save(full_path)
        print(f'\nFile "{full_path}" saved successfully!')
        
        return choice, full_path
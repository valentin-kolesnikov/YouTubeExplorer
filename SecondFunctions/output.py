import sys
from datetime import datetime
from pathlib import Path

from docx import Document


def output_videos(results, statrequest, keywords=None, one_video_info=False):
    parsed_videos = []

    if not one_video_info:
        print(f"Your request: {keywords}")

    for number, item in enumerate(statrequest["items"], start=1):
        title = item["snippet"]["title"]
        video_id = item["id"]
        published_at = item["snippet"]["publishedAt"]
        dt = datetime.fromisoformat(published_at)
        UTC = dt.strftime("%z")
        formatted_date = dt.strftime("%d.%m.%Y %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

        likes = item["statistics"].get("likeCount", "No")
        dislikes = results.get(video_id, {}).get("dislikes", "No")
        views = item["statistics"].get("viewCount", "No")
        comments = item["statistics"].get("commentCount", "No")

        channelName = item["snippet"]["channelTitle"]
        channelId = item["snippet"]["channelId"]

        description = None
        if one_video_info:
            description = item["snippet"]["description"]

        parsed_videos.append({
            "title": title,
            "video_id": video_id,
            "formatted_date": formatted_date,
            "likes": likes,
            "dislikes": dislikes,
            "views": views,
            "comments": comments,
            "channelName": channelName,
            "channelId": channelId,
            "description": description
        })


        if not one_video_info:
            print(f"\n{number}.")
        print(
            f"Title: {title}\n"
            f"Video Link: https://www.youtube.com/watch?v={video_id}\n"
            f"Views: {views} | Likes: {likes} | Dislikes: {dislikes} | Comments: {comments}\n"
            f"Date: {formatted_date}\n"
            f"Channel: {channelName}\n"
            f"Channel URL: https://www.youtube.com/channel/{channelId}")
        if one_video_info:
            print(f"\nDescription:\n=================================\n{description}\n=================================")

    return parsed_videos



def save_docx(parsed_videos, keywords, region=None):
    choice = input("\n\nDo you want to save videos in a DOCX file? (y/n): ").strip().lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").strip().lower()
        
    if choice == "n":
        return choice, None

    elif choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = Path(app_folder, "Videos_in_DOCX")
        youtube_folder.mkdir(parents=True, exist_ok=True)

        full_path = Path(youtube_folder, f"{keywords}.docx")

        counter = 0
        while full_path.exists():
            counter += 1
            full_path = Path(youtube_folder, f"{keywords} ({counter}).docx")

        doc = Document()
        doc.add_heading("Video Report for YouTube", 0)


        doc.add_heading(f"Request: {keywords}", level=1)

        if region:
            doc.add_heading(f"Region: {region}", level=2)
        
        doc.add_heading("The list of videos: ", level=1)

        for number, info in enumerate(parsed_videos, start=1):
            if info["description"]:
                doc.add_heading(f"{info['title']}", level=2)

            else:
                doc.add_heading(f"\n{number}. {info['title']}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Video Link: ").bold = True
            p.add_run(f"https://www.youtube.com/watch?v={info['video_id']}\n")
            
            p.add_run("Statistics: ").bold = True
            p.add_run(f"{info['views']} views; {info['likes']} likes; {info['dislikes']} dislikes; {info['comments']} comments\n")
            
            p.add_run("Date: ").bold = True
            p.add_run(f"{info['formatted_date']}\n")
            
            p.add_run("Channel: ").bold = True
            p.add_run(f"{info['channelName']}\n")
            
            p.add_run("Channel URL: ").bold = True
            p.add_run(f"https://www.youtube.com/channel/{info['channelId']}")

            if info['description']:
                p.add_run("Channel URL: ").bold = True
                p.add_run(f"\nDescription:\n=================================\n{info['description']}\n=================================")
        

        doc.save(full_path)
        print(f'\nFile "{full_path}" saved successfully!')
        
        return choice, full_path

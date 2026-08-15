import sys
from datetime import datetime
from pathlib import Path

from docx import Document


def output_channel_info(result, statrequests, get_answers, snistics, keywords=None):

    print(f"Channel: {snistics["title"]}\n"
          f"https://www.youtube.com/channel/{snistics["channelId"]}\n"
          f"CustomUrl: {snistics["customUrl"]}\n"
          f"{snistics["subscriberCount"]} subs; {snistics["videoCount"]} videos; {snistics["viewCount"]} views.\n"
          f"Registration date: {snistics["publishedAt"]}\n"
          f"\nDescription:\n=================================\n{snistics["description"]}\n=================================\n")
    
    print("-" * 50)

    if get_answers == "y":
        print(f"Your request: {keywords}\n")
        print("Your received videos:")
        
    elif get_answers == "n":
        print("Four videos from the newest:")

    parsed_videos = []

    for number, item in enumerate(statrequests["items"], start=1):
        title_video = item["snippet"]["title"]
        video_id = item["id"]
        published_at = item["snippet"]["publishedAt"]
        dt = datetime.fromisoformat(published_at.replace("Z", "+00:00"))
        formatted_date = dt.strftime("%d.%m.%Y %H:%M:%S")

        ch_likes = item["statistics"].get("likeCount", "No")
        ch_dislikes = result.get(video_id, {}).get("dislikes", "No")
        ch_views = item["statistics"].get("viewCount", "No")
        ch_comments = item["statistics"].get("commentCount", "No")

        parsed_videos.append({
                    "title": title_video,
                    "video_id": video_id,
                    "formatted_date": formatted_date,
                    "likes": ch_likes,
                    "dislikes": ch_dislikes,
                    "views": ch_views,
                    "comments": ch_comments,
                })

        print(
            f"\n{number}.\n"
            f"{title_video}\n"
            f"Video Link: https://www.youtube.com/watch?v={video_id}\n"
            f"{ch_views} views; {ch_likes} likes; {ch_dislikes} dislikes; {ch_comments} comments\n"
            f"Date: {formatted_date}")

    return parsed_videos







def save_docx(snistics, parsed_videos, keywords=None):
    choice = input("\n\nDo you want to save channel's info in a DOCX file? (y/n): ").strip().lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").strip().lower()
        
    if choice == "n":
        return choice, None

    elif choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = Path(app_folder, "Channels_in_DOCX")
        youtube_folder.mkdir(parents=True, exist_ok=True)

        full_path = Path(youtube_folder, f"{snistics["title"]}.docx")

        counter = 0
        while full_path.exists():
            counter += 1
            full_path = Path(youtube_folder, f"{snistics["title"]} ({counter}).docx")

        doc = Document()
        doc.add_heading("Channel Report for YouTube", 0)

        doc.add_heading("Channel Information", level=1)
        
        p_channel = doc.add_paragraph()
        p_channel.add_run("Title: ").bold = True
        p_channel.add_run(f"{snistics['title']}\n")
        
        p_channel.add_run("Channel Link: ").bold = True
        p_channel.add_run(f"https://www.youtube.com/channel/{snistics['channelId']}\n")
        
        p_channel.add_run("Custom URL: ").bold = True
        p_channel.add_run(f"{snistics.get('customUrl', 'No')}\n")
        
        p_channel.add_run("Statistics: ").bold = True
        p_channel.add_run(f"{snistics['subscriberCount']} subs; {snistics['videoCount']} videos; {snistics['viewCount']} views\n")
        
        p_channel.add_run("Registration Date: ").bold = True
        p_channel.add_run(f"{snistics['publishedAt']}\n")

        doc.add_heading("Description:", level=2)
        doc.add_paragraph(snistics['description'])

        if keywords:
            doc.add_heading(f"Request: {keywords}", level=1)

        doc.add_heading("The list of videos: ", level=1)

        for number, info in enumerate(parsed_videos, start=1):
            doc.add_heading(f"{number}. {info['title']}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Video Link: ").bold = True
            p.add_run(f"https://www.youtube.com/watch?v={info['video_id']}\n")
            
            p.add_run("Statistics: ").bold = True
            p.add_run(f"{info['views']} views; {info['likes']} likes; {info['dislikes']} dislikes; {info['comments']} comments\n")
            
            p.add_run("Date: ").bold = True
            p.add_run(f"{info['formatted_date']}\n")
        

        doc.save(full_path)
        print(f'\nFile "{full_path}" saved successfully!')
        
        return choice, full_path
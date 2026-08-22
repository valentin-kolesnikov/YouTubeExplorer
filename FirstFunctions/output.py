import sys
from datetime import datetime
from pathlib import Path

from docx import Document

from Patterns.save_history import clear


def count_keys(comments, search_terms):
    counts = {key_word: 0 for key_word in search_terms}

    amount_comments = len(comments)
    amount_replies = sum(len(comment["replies"]) for comment in comments)
    total_comments = amount_comments + amount_replies

    for comment in comments:
        for key_word in search_terms:
            counts[key_word] += comment.lower().count(key_word.lower())

    print(f"Total comments and replies: {total_comments}\n")
    print(f"Total comments: {amount_comments}")
    print(f"Total replies: {amount_replies}\n")

    for key_word, count in counts.items():
        print(f"{key_word}: {count}\n")

    return amount_comments, amount_replies, total_comments, counts


    


def number_comments(comments, channel_id, channel_title):
    number = input("How many comments do you need? (Only the number of comments is considered): ").strip()
    while not number.isdigit():
        number = input("\nEnter again: ").strip()

    clear()

    print(f"Channel: {channel_title}")
    print(f"\nThe channel URL: https://www.youtube.com/channel/{channel_id}")

    for number, comment in enumerate(comments, start=1):
        
        top_published_at = comment["published_at"]
        dt1 = datetime.fromisoformat(top_published_at)
        UTC1 = dt1.strftime("%z")
        top_formatted_date = dt1.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC1[:3]}:{UTC1[3:]})"

        top_updated_at = comment["updated_at"]
        dt = datetime.fromisoformat(top_updated_at)
        UTC = dt.strftime("%z")
        top_updated_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

        print(f"\n{number}."
              f"\nComment: {comment["text"]}"
              f"\nAuthor: {comment["author"]}"
              f"\nLike Count: {comment["likeCount"]}"
              f"\nPublished at: {top_formatted_date}"
              f"\nUpdated at: {top_updated_formatted_date}"
              f"\nComment ID: {comment["id"]}")
              

        for number_reply, reply in enumerate(comment["replies"], start=1):

            reply_published_at = reply["published_at"]
            dt = datetime.fromisoformat(reply_published_at)
            UTC = dt.strftime("%z")
            reply_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

            reply_updated_at = reply["updated_at"]
            dt = datetime.fromisoformat(reply_updated_at)
            UTC = dt.strftime("%z")
            reply_updated_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

            print(f"\n\t{number}.{number_reply}."
                  f"\n\tReply: {reply["text"]}"
                  f"\n\tAuthor: {reply["author"]}"
                  f"\n\tLike Count: {reply["likeCount"]}"
                  f"\n\tPublished at: {reply_formatted_date}"
                  f"\n\tUpdated at: {reply_updated_formatted_date}"
                  f"\n\tReply ID: {reply["id"]}")










def save_docx(comments, channel_id, channel_title, counts, amount_comments, video_id, amount_replies, total_comments):
    choice = input("\n\nDo you want to save the comments in a DOCX file? (y/n): ").strip().lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").strip().lower()
        
    if choice == "n":
        return choice, None

    elif choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = Path(app_folder, "Comments_in_DOCX")
        youtube_folder.mkdir(parents=True, exist_ok=True)

        full_path = Path(youtube_folder, f"{video_id}.docx")

        counter = 0
        while full_path.exists():
            counter += 1
            full_path = Path(youtube_folder, f"{video_id} ({counter}).docx")
            

        doc = Document()
        doc.add_heading("Comment Report for YouTube", 0)

        doc.add_paragraph(f"Channel: {channel_title}")
        doc.add_paragraph(f"Channel link: https://www.youtube.com/channel/{channel_id}")
        doc.add_paragraph(f"Video link: https://www.youtube.com/watch?v={video_id}")

        doc.add_paragraph(f"Total comments and replies found: {total_comments}")
        doc.add_paragraph(f"Total comments found: {amount_comments}")
        doc.add_paragraph(f"Total replies found: {amount_replies}")

        doc.add_heading("Keyword Statistics:", level=1)
        if counts == {}:
            doc.add_paragraph("None")
        else:
            for key_word, count in counts.items():
                doc.add_paragraph(f"{key_word}: {count}", style='List Bullet')
            
        doc.add_heading("Selected Comments:", level=1)
        for index, comment in enumerate(comments, start=1):

            top_published_at = comment["published_at"]
            dt1 = datetime.fromisoformat(top_published_at)
            UTC1 = dt1.strftime("%z")
            top_formatted_date = dt1.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC1[:3]}:{UTC1[3:]})"
    
            top_updated_at = comment["updated_at"]
            dt = datetime.fromisoformat(top_updated_at)
            UTC = dt.strftime("%z")
            top_updated_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

            doc.add_paragraph(f"\n{index}."
                                f"\nComment: {comment['text']}"
                                f"\nAuthor: {comment['author']}"
                                f"\nLike Count: {comment['likeCount']}"
                                f"\nPublished at: {top_formatted_date}"
                                f"\nUpdated at: {top_updated_formatted_date}"
                                f"\nComment ID: {comment['id']}")
            doc.add_paragraph("-" * 20)

            for index_reply, reply in enumerate(comment["replies"], start=1):
                reply_published_at = reply["published_at"]
                dt = datetime.fromisoformat(reply_published_at)
                UTC = dt.strftime("%z")
                reply_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

                reply_updated_at = reply["updated_at"]
                dt = datetime.fromisoformat(reply_updated_at)
                UTC = dt.strftime("%z")
                reply_updated_formatted_date = dt.strftime("%Y.%m.%d %H:%M:%S") + f" (UTC{UTC[:3]}:{UTC[3:]})"

                doc.add_paragraph(f"\n\t{index}.{index_reply}."
                                    f"\n\tReply: {reply['text']}"
                                    f"\n\tAuthor: {reply['author']}"
                                    f"\n\tLike Count: {reply['likeCount']}"
                                    f"\n\tPublished at: {reply_formatted_date}"
                                    f"\n\tUpdated at: {reply_updated_formatted_date}"
                                    f"\n\tReply ID: {reply['id']}")
                doc.add_paragraph("-" * 20)
            

        doc.save(full_path)
        print(f'\nFile "{full_path}" saved successfully!')
        
        return choice, full_path